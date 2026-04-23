"""
Parse uploaded report files (PDF / DOCX) into text + tables.

We extract text page-by-page and render any detected tables as Markdown so
they survive the round-trip into an LLM prompt. Images and embedded figures
are NOT processed in this release (text + tables only, per user spec).
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import Any, Optional

# Soft imports — availability is reported via parser_availability() so the
# web server can boot with a clear error message rather than crash at import.
try:
    import fitz  # type: ignore  # PyMuPDF
    _HAS_PYMUPDF = True
except Exception:
    _HAS_PYMUPDF = False

try:
    import docx  # type: ignore  # python-docx
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


SUPPORTED_EXT = {".pdf", ".docx"}


@dataclass
class ParseResult:
    """Result of parsing one report file."""
    ext: str                                # ".pdf" | ".docx"
    page_count: int                         # pages for PDF, 1 for DOCX
    text: str                               # full concatenated text
    tables_markdown: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def text_length(self) -> int:
        return len(self.text)

    @property
    def tables_count(self) -> int:
        return len(self.tables_markdown)

    def to_prompt_block(self, max_chars: int = 60_000) -> str:
        """
        Serialize the parsed report for injection into the system prompt.
        If the content exceeds max_chars, a tail-truncation notice is added.
        """
        parts: list[str] = []
        if self.text.strip():
            parts.append("## 报表正文\n\n" + self.text.strip())
        if self.tables_markdown:
            parts.append("## 报表表格 (共 {} 张)\n\n".format(len(self.tables_markdown))
                         + "\n\n".join(self.tables_markdown))
        blob = "\n\n".join(parts) if parts else "(报表内容为空)"
        if len(blob) > max_chars:
            kept = blob[:max_chars]
            omitted = len(blob) - max_chars
            blob = kept + f"\n\n... [已截断,省略 {omitted} 个字符]"
        return blob


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parser_availability() -> dict[str, bool]:
    """Report which backends are importable — used by /api/meta."""
    return {"pdf": _HAS_PYMUPDF, "docx": _HAS_DOCX}


def parse_report(file_path: str) -> ParseResult:
    """Detect by extension and dispatch to the right parser."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        if not _HAS_PYMUPDF:
            raise RuntimeError(
                "PDF 解析需要 pymupdf,请先 `pip install pymupdf`(或安装 web extras)。"
            )
        return _parse_pdf(file_path)
    if ext == ".docx":
        if not _HAS_DOCX:
            raise RuntimeError(
                "Word 解析需要 python-docx,请先 `pip install python-docx`(或安装 web extras)。"
            )
        return _parse_docx(file_path)
    raise ValueError(f"不支持的文件类型: {ext}  (仅支持 .pdf / .docx)")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(path: str) -> ParseResult:
    text_parts: list[str] = []
    tables_md: list[str] = []
    warnings: list[str] = []

    with fitz.open(path) as doc:  # type: ignore[union-attr]
        page_count = doc.page_count
        for idx, page in enumerate(doc, start=1):
            # Text
            try:
                page_text = page.get_text("text") or ""
            except Exception as e:
                page_text = ""
                warnings.append(f"第 {idx} 页文本抽取失败: {e}")
            if page_text.strip():
                text_parts.append(f"--- Page {idx} ---\n{page_text.rstrip()}")

            # Tables — PyMuPDF >= 1.23 provides page.find_tables()
            finder = getattr(page, "find_tables", None)
            if callable(finder):
                try:
                    tf = finder()
                    tables = getattr(tf, "tables", []) or []
                    for t_idx, tbl in enumerate(tables, start=1):
                        rows = _extract_table_rows(tbl)
                        md = _rows_to_markdown(rows)
                        if md:
                            tables_md.append(f"### Page {idx} · Table {t_idx}\n\n{md}")
                except Exception as e:
                    warnings.append(f"第 {idx} 页表格抽取失败: {e}")

    return ParseResult(
        ext=".pdf",
        page_count=page_count,
        text="\n\n".join(text_parts),
        tables_markdown=tables_md,
        warnings=warnings,
    )


def _extract_table_rows(table_obj: Any) -> list[list[str]]:
    """Normalize PyMuPDF Table.extract() result into a list of string rows."""
    try:
        raw = table_obj.extract()
    except Exception:
        return []
    if not raw:
        return []
    norm: list[list[str]] = []
    for row in raw:
        norm.append(["" if cell is None else str(cell).strip() for cell in row])
    return norm


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parse_docx(path: str) -> ParseResult:
    document = docx.Document(path)  # type: ignore[union-attr]
    text_parts: list[str] = []
    tables_md: list[str] = []
    warnings: list[str] = []

    # Paragraphs
    for para in document.paragraphs:
        line = para.text.rstrip()
        if line:
            text_parts.append(line)

    # Tables
    for t_idx, tbl in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in tbl.rows:
            rows.append([cell.text.replace("\n", " ").strip() for cell in row.cells])
        md = _rows_to_markdown(rows)
        if md:
            tables_md.append(f"### Table {t_idx}\n\n{md}")

    return ParseResult(
        ext=".docx",
        page_count=1,   # Word documents are flow-based; no true page count
        text="\n".join(text_parts),
        tables_markdown=tables_md,
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def _rows_to_markdown(rows: list[list[str]]) -> str:
    """Render a rectangular table as a Markdown table. Returns "" if empty."""
    if not rows:
        return ""
    # Normalize to the widest row
    width = max(len(r) for r in rows)
    if width == 0:
        return ""
    padded = [r + [""] * (width - len(r)) for r in rows]
    header, *body = padded
    header = [_cell_clean(c) or f"col{i+1}" for i, c in enumerate(header)]

    lines: list[str] = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in body:
        lines.append("| " + " | ".join(_cell_clean(c) for c in row) + " |")
    return "\n".join(lines)


def _cell_clean(cell: Optional[str]) -> str:
    if cell is None:
        return ""
    # Strip newlines and pipe so the Markdown table survives
    return str(cell).replace("\n", " ").replace("|", "/").strip()
